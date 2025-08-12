import yfinance as yf
from docx import Document
import re
from datetime import datetime

# Load the Word document (replace with your actual file name)
doc = Document("my_equity_brief.docx")

# Extract full text
full_text = "\n".join([p.text for p in doc.paragraphs])

# Find ticker symbols (simple pattern for uppercase tickers)
tickers = re.findall(r'\b[A-Z]{1,5}\b', full_text)
tickers = list(set(tickers))  # remove duplicates

# Fetch latest data for each ticker
data_map = {}
for t in tickers:
    try:
        stock = yf.Ticker(t)
        info = stock.info
        data_map[t] = {
            "price": info.get("currentPrice"),
            "pe": info.get("trailingPE"),
            "marketCap": info.get("marketCap"),
            "eps": info.get("trailingEps")
        }
    except Exception as e:
        print(f"Error fetching {t}: {e}")

# Update paragraphs with latest prices
for p in doc.paragraphs:
    for t, values in data_map.items():
        if values["price"]:
            p.text = re.sub(
                rf'({t}.*?\$?\d+(\.\d+)?)',
                f"{t} latest price: ${values['price']}",
                p.text
            )

# Append investment recommendation summary
doc.add_paragraph("\n---\nUpdated on " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
for t, values in data_map.items():
    rec = "HOLD"
    score = 0
    if values["eps"] and values["eps"] > 0: score += 1
    if values["pe"] and values["pe"] < 20: score += 1
    if score >= 2:
        rec = "BUY"
    elif score <= 0:
        rec = "SELL"
    doc.add_paragraph(f"{t}: Recommendation → {rec} (Price: ${values['price']}, PE: {values['pe']})")

# Save updated document
doc.save("my_equity_brief_UPDATED.docx")

print("Update complete. Check my_equity_brief_UPDATED.docx")

